import streamlit as st
import yfinance as yf
import pandas as pd
import numpy as np
from scipy.stats import norm
import plotly.graph_objects as go
import matplotlib.pyplot as plt
from datetime import datetime, timedelta
import time
import requests
import xml.etree.ElementTree as ET
from io import BytesIO
import google.generativeai as genai
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# --- PAGE CONFIGURATION ---
st.set_page_config(page_title="Options Command Center", layout="wide", page_icon="🚀")

# --- PASSWORD PROTECTION ---
def check_password():
    if "password_correct" not in st.session_state:
        st.session_state["password_correct"] = False
    
    def password_entered():
        if st.session_state["password"] == st.secrets["passwords"]["main_password"]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if not st.session_state["password_correct"]:
        st.text_input("🔑 Enter Password", type="password", on_change=password_entered, key="password")
        return False
    return True

if not check_password():
    st.stop()

# --- HELPER FUNCTIONS ---

def get_google_news(ticker):
    try:
        url = f"https://news.google.com/rss/search?q={ticker}+stock&hl=en-US&gl=US&ceid=US:en"
        response = requests.get(url, timeout=5)
        root = ET.fromstring(response.content)
        news_items = []
        for item in root.findall('.//item')[:5]:
            title = item.find('title').text
            link = item.find('link').text
            pubDate = item.find('pubDate').text
            source = "Google News"
            if " - " in title:
                parts = title.rsplit(" - ", 1)
                title = parts[0]
                source = parts[1]
            news_items.append({'title': title, 'link': link, 'publisher': source, 'published': pubDate})
        return news_items
    except Exception as e:
        return []

def fetch_with_retry(func, *args, retries=3):
    for i in range(retries):
        try:
            return func(*args)
        except Exception as e:
            error_msg = str(e).lower()
            if "too many requests" in error_msg or "429" in error_msg:
                wait_time = 10 * (i + 1)
                time.sleep(wait_time)
                continue
            raise e
    return func(*args)

@st.cache_data(ttl=3600) 
def get_stock_history_and_info(ticker_symbol):
    def _get():
        stock = yf.Ticker(ticker_symbol)
        history = stock.history(period="3mo")
        info = stock.info
        news = get_google_news(ticker_symbol)
        return history, info, news
    return fetch_with_retry(_get)

@st.cache_data(ttl=3600)
def get_option_chain_data(ticker_symbol, date):
    def _get():
        stock = yf.Ticker(ticker_symbol)
        opt_chain = stock.option_chain(date)
        calls = opt_chain.calls
        calls['type'] = 'call'
        puts = opt_chain.puts
        puts['type'] = 'put'
        full = pd.concat([calls, puts])
        return full, calls, puts
    return fetch_with_retry(_get)

def get_ticker_object(ticker_symbol):
    return yf.Ticker(ticker_symbol)

def black_scholes_price(S, K, T, r, sigma, option_type='call'):
    if T <= 0: return max(0, S - K) if option_type == 'call' else max(0, K - S)
    d1 = (np.log(S / K) + (r + 0.5 * sigma ** 2) * T) / (sigma * np.sqrt(T))
    d2 = d1 - sigma * np.sqrt(T)
    if option_type == 'call':
        price = S * norm.cdf(d1) - K * np.exp(-r * T) * norm.cdf(d2)
    else:
        price = K * np.exp(-r * T) * norm.cdf(-d2) - S * norm.cdf(-d1)
    return price

def calculate_greeks(S, K, T, r, sigma, option_type='call'):
    if T <= 0 or sigma <= 0: return 0, 0, 0
    d1 = (np.log(S / K) + (r + 0.5 * sigma ** 2) * T) / (sigma * np.sqrt(T))
    d2 = d1 - sigma * np.sqrt(T)
    if option_type == 'call':
        delta = norm.cdf(d1)
    else:
        delta = norm.cdf(d1) - 1 
    gamma = norm.pdf(d1) / (S * sigma * np.sqrt(T))
    term1 = -(S * sigma * norm.pdf(d1)) / (2 * np.sqrt(T))
    term2 = r * K * np.exp(-r * T) * norm.cdf(d2)
    if option_type == 'call':
        theta_annual = term1 - term2
    else:
        theta_annual = term1 + r * K * np.exp(-r * T) * norm.cdf(-d2)
    theta_daily = theta_annual / 365.0
    return delta, gamma, theta_daily

def calculate_max_pain(options_chain):
    strikes = options_chain['strike'].unique()
    max_pain_data = []
    for strike in strikes:
        calls_at_strike = options_chain[options_chain['type'] == 'call']
        puts_at_strike = options_chain[options_chain['type'] == 'put']
        call_loss = calls_at_strike.apply(lambda x: max(0, strike - x['strike']) * x['openInterest'], axis=1).sum()
        put_loss = puts_at_strike.apply(lambda x: max(0, x['strike'] - strike) * x['openInterest'], axis=1).sum()
        max_pain_data.append({'strike': strike, 'total_loss': call_loss + put_loss})
    df_pain = pd.DataFrame(max_pain_data)
    if df_pain.empty: return 0
    return df_pain.loc[df_pain['total_loss'].idxmin()]['strike']

# --- PORTFOLIO FUNCTIONS ---
def recalculate_portfolio(df):
    if df.empty: return df
    df['Qty'] = pd.to_numeric(df['Qty'], errors='coerce').fillna(0)
    df['Bought Price'] = pd.to_numeric(df['Bought Price'], errors='coerce').fillna(0.0)
    df['Current Price'] = pd.to_numeric(df['Current Price'], errors='coerce').fillna(0.0)
    multipliers = np.where(df['Type'] == 'Option', 100, 1)
    df['Total Cost'] = df['Bought Price'] * df['Qty'] * multipliers
    df['Current Value'] = df['Current Price'] * df['Qty'] * multipliers
    df['Total P/L'] = df['Current Value'] - df['Total Cost']
    return df

# --- PLOTS ---
def plot_greeks_interactive(current_price, strike, days_left, iv, opt_type):
    prices = np.linspace(strike * 0.8, strike * 1.2, 100)
    T = max(days_left / 365.0, 0.001)
    deltas = [calculate_greeks(p, strike, T, 0.045, iv, opt_type)[0] for p in prices]
    gammas = [calculate_greeks(p, strike, T, 0.045, iv, opt_type)[1] for p in prices]
    curr_d, _, _ = calculate_greeks(current_price, strike, T, 0.045, iv, opt_type)
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=prices, y=deltas, mode='lines', name='Delta', line=dict(color='#4DA6FF', width=3)))
    fig.add_trace(go.Scatter(x=prices, y=gammas, mode='lines', name='Gamma', line=dict(color='#00FF7F', width=2, dash='dash'), yaxis="y2"))
    fig.add_trace(go.Scatter(x=[current_price], y=[curr_d], mode='markers', name='You', marker=dict(color='white', size=10)))
    fig.update_layout(title=f"Greeks Profile ({opt_type.upper()})", template="plotly_dark", height=450, yaxis2=dict(overlaying="y", side="right"), dragmode='pan')
    return fig

def plot_simulation_interactive(S, K, days_left, iv, opt_type, r=0.045, purchase_price=0):
    prices = np.linspace(S * 0.8, S * 1.2, 100)
    T1 = max(days_left / 365.0, 0.0001)
    pnl_today = [black_scholes_price(p, K, T1, r, iv, opt_type) - purchase_price for p in prices]
    T2 = max((days_left / 2) / 365.0, 0.0001)
    pnl_half = [black_scholes_price(p, K, T2, r, iv, opt_type) - purchase_price for p in prices]
    T3 = 0.0001
    pnl_exp = [black_scholes_price(p, K, T3, r, iv, opt_type) - purchase_price for p in prices]
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=prices, y=pnl_today, mode='lines', name='Today (T+0)', line=dict(color='#4DA6FF', width=3)))
    fig.add_trace(go.Scatter(x=prices, y=pnl_half, mode='lines', name=f'Halfway', line=dict(color='#FFD700', width=2, dash='dash')))
    fig.add_trace(go.Scatter(x=prices, y=pnl_exp, mode='lines', name='Expiration', line=dict(color='#FF4B4B', width=2, dash='dot')))
    fig.add_hline(y=0, line_color="white", opacity=0.5)
    fig.add_vline(x=S, line_color="gray", line_dash="dash", annotation_text="Current Price")
    fig.update_layout(title=f"🔮 Future P&L Simulator ({opt_type.upper()})", xaxis_title="Price", yaxis_title="P&L", template="plotly_dark", height=500, dragmode='pan')
    return fig

def plot_whale_activity_interactive(df, current_strike, opt_type):
    strikes = sorted(df['strike'].unique())
    try: idx = strikes.index(current_strike)
    except: idx = 0
    start = max(0, idx - 4); end = min(len(strikes), idx + 5)
    subset = df[df['strike'].isin(strikes[start:end])]
    color = '#FF4B4B' if opt_type == 'put' else '#00FF7F' 
    fig = go.Figure()
    fig.add_trace(go.Bar(x=subset['strike'], y=subset['openInterest'], name='OI', marker_color='#4DA6FF'))
    fig.add_trace(go.Bar(x=subset['strike'], y=subset['volume'], name='Vol', marker_color=color))
    fig.update_layout(title=f"Whale Detector ({opt_type.upper()})", template="plotly_dark", height=450, dragmode='pan')
    return fig

def plot_flow_battle_interactive(calls, puts, current_strike):
    c = calls.groupby('strike')['volume'].sum(); p = puts.groupby('strike')['volume'].sum()
    df = pd.merge(c, p, on='strike', how='outer').fillna(0).sort_index()
    try: idx = df.index.get_loc(current_strike)
    except: idx = len(df)//2
    sub = df.iloc[max(0, idx-5):min(len(df), idx+6)]
    fig = go.Figure()
    fig.add_trace(go.Bar(x=sub.index, y=sub['volume_x'], name='Bulls (Calls)', marker_color='#00FF7F'))
    fig.add_trace(go.Bar(x=sub.index, y=sub['volume_y'], name='Bears (Puts)', marker_color='#FF4B4B'))
    fig.update_layout(title="Battle Map", template="plotly_dark", height=450, dragmode='pan')
    return fig

def add_colored_box(doc, text, color_hex):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = text
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generate_full_dossier(data):
    doc = Document()
    head = doc.add_heading(f"MISSION REPORT: {data['ticker']} ({data['type'].upper()})", 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_heading("1. Executive Summary", 1)
    t = doc.add_table(rows=2, cols=4); t.style = 'Table Grid'
    r = t.rows[1].cells
    r[0].text = data['ticker']; r[1].text = f"${data['price']:.2f}"
    r[2].text = f"${data['strike']} {data['type'].upper()}"; r[3].text = str(data['exp'])
    doc.add_heading("2. Key Intel Metrics", 1)
    p = doc.add_paragraph()
    p.add_run(f"• IV: ").bold = True; p.add_run(f"{data['iv']*100:.2f}%\n")
    p.add_run(f"• Rule of 16: ").bold = True; p.add_run(f"${data['daily_move']:.2f}\n")
    p.add_run(f"• Volume: ").bold = True; p.add_run(f"{data['volume']:,}")
    b = BytesIO(); doc.save(b); return b

# --- UPGRADED SCANNER (Targeting Monthly/Feb Expirations) ---
def run_scan(tickers):
    res = []
    bar = st.progress(0); txt = st.empty()
    try: batch_data = yf.download(tickers, period="1d", group_by='ticker', progress=False)
    except: batch_data = pd.DataFrame()
    
    for i, t in enumerate(tickers):
        time.sleep(1.0) # Faster scan
        try:
            txt.text(f"Scanning options for {t}...")
            curr = 0
            if not batch_data.empty:
                try:
                    if len(tickers) > 1: curr = batch_data[t]['Close'].iloc[-1]
                    else: curr = batch_data['Close'].iloc[-1]
                except: pass
            
            stk = yf.Ticker(t)
            if curr == 0: curr = stk.history(period='1d')['Close'].iloc[-1]
            dates = stk.options
            if not dates: continue
            
            # --- LOGIC CHANGE HERE ---
            # Old Logic: target_date = dates[0] (Next week)
            # New Logic: Find the first date that is at least 25 days away (Feb Monthly)
            target_date = dates[0]
            for d in dates:
                dt_obj = datetime.strptime(d, "%Y-%m-%d")
                days_out = (dt_obj - datetime.now()).days
                if days_out >= 25: # Looks for trades ~1 month out (Feb 20)
                    target_date = d
                    break
            # -------------------------

            calls = stk.option_chain(target_date).calls
            calls['diff'] = abs(calls['strike'] - curr)
            atm = calls.loc[calls['diff'].idxmin()]
            iv = atm['impliedVolatility']
            days_to_exp = (datetime.strptime(target_date, "%Y-%m-%d") - datetime.now()).days
            if days_to_exp < 1: days_to_exp = 1
            _, gamma, _ = calculate_greeks(curr, atm['strike'], days_to_exp/365, 0.045, iv)
            vol_oi = atm['volume'] / atm['openInterest'] if atm['openInterest'] > 0 else 0
            moneyness = ((curr - atm['strike']) / atm['strike']) * 100
            
            res.append({
                'Ticker': t, 
                'ATM Strike': atm['strike'], 
                'Exp Date': target_date, # Now shows Feb 20
                'Price': atm['lastPrice'], 
                'Vol': atm['volume'], 
                'Vol/OI': round(vol_oi, 2), 
                'Money%': round(moneyness, 2), 
                'Gamma': round(gamma, 4)
            })
        except: pass
        bar.progress((i+1)/len(tickers))
    txt.empty(); bar.empty()
    return pd.DataFrame(res)
   
# --- INITIALIZE SESSION STATE ---
if "portfolio" not in st.session_state:
    st.session_state["portfolio"] = pd.DataFrame(columns=[
        "Ticker", "Type", "Details", "Qty", "Bought Price", "Current Price", "Total Cost", "Current Value", "Total P/L"
    ])

# --- MAIN NAVIGATION ---
st.sidebar.markdown("## 🧭 Navigation")
app_mode = st.sidebar.radio("Go To:", ["Analysis Dashboard 📊", "Portfolio Tracker 📒"])
st.sidebar.markdown("---")

# =========================================================
# PAGE 1: ANALYSIS DASHBOARD (Original Features)
# =========================================================
if app_mode == "Analysis Dashboard 📊":
    st.sidebar.markdown("## ⚙️ Dashboard Settings")
    ticker = st.sidebar.text_input("Ticker Symbol", value="PEP").upper()
    option_type = st.sidebar.selectbox("Option Type", ["Call", "Put"]).lower()
    strike_price = st.sidebar.number_input("Strike Price ($)", value=148.0)

    if st.sidebar.button("🔄 Force Refresh Data"):
        st.cache_data.clear(); st.rerun()

    if ticker:
        try:
            stock_conn = get_ticker_object(ticker)
            with st.spinner('Fetching market data...'):
                history, info, news_data = get_stock_history_and_info(ticker)
                current_price = info.get('currentPrice', history['Close'].iloc[-1])
                prev_close = info.get('previousClose', history['Close'].iloc[-2])
                
                expirations = stock_conn.options
                if not expirations: st.error("No options data."); st.stop()
                selected_date = st.sidebar.selectbox("Expiration Date", expirations)
                
                full_chain, calls, puts = get_option_chain_data(ticker, selected_date)
                active_chain = calls if option_type == 'call' else puts
                
                specific_contract = active_chain.iloc[(active_chain['strike'] - strike_price).abs().argsort()[:1]]
                contract_iv = specific_contract.iloc[0]['impliedVolatility']
                
                days_left = (datetime.strptime(selected_date, "%Y-%m-%d") - datetime.now()).days
                if days_left < 1: days_left = 1
                
                theo_price = black_scholes_price(current_price, strike_price, days_left/365, 0.045, contract_iv, option_type)
                d, g, t = calculate_greeks(current_price, strike_price, days_left/365, 0.045, contract_iv, option_type)
                max_pain_val = calculate_max_pain(full_chain)
                daily_move = (contract_iv * 100 / 16) / 100 * current_price

            st.title(f"📊 {ticker} {option_type.upper()} Command Center")
            col1, col2, col3 = st.columns(3)
            col1.metric("Current Price", f"${current_price:.2f}", f"{current_price - prev_close:.2f}")
            col2.metric(f"Your {option_type.title()}", f"${strike_price:.2f}")
            col3.metric("Selected Expiration", selected_date)
            st.markdown("---")

            tabs = st.tabs([
                "1. Price", "2. Volume", "3. IV", "4. Rule of 16", 
                "5. Whale Detector", "6. Risk & Profit", "7. Max Pain", "8. News", 
                "9. 🤖 Chart Analyst", "10. 💬 Strategy Engine", "11. 🔮 Future Simulator", "12. 🌊 Flow Monitor",
                "13. 🔍 ATM Scanner"
            ])

            with tabs[0]: 
                st.subheader(f"1. {ticker} Stock Price")
                st.line_chart(history['Close'])
                # REMOVED THE "ESTIMATED OPTION PRICE" CHART HERE
                st.info("ℹ️ Note: Historical option price charts (candlesticks) are not available on free data sources. Please check your brokerage for the official trade history.")

            with tabs[1]: st.metric("Volume", f"{info.get('volume', 0):,}")
            with tabs[2]: st.metric("IV", f"{contract_iv * 100:.2f}%")
            with tabs[3]: st.metric("Expected Daily Move", f"${daily_move:.2f}")

            with tabs[4]: # Whale
                st.header(f"Whale Detector ({option_type.upper()})")
                st.plotly_chart(plot_whale_activity_interactive(active_chain, strike_price, option_type), use_container_width=True)

            with tabs[5]: # Greeks & Calc
                st.header("Risk & Profit Hub")
                c1, c2, c3 = st.columns(3)
                c1.metric("Delta", f"{d:.2f}"); c2.metric("Gamma", f"{g:.3f}"); c3.metric("Theta", f"{t:.3f}")
                st.plotly_chart(plot_greeks_interactive(current_price, strike_price, days_left, contract_iv, option_type), use_container_width=True)
                st.markdown("---")
                st.subheader("🎯 Profit Target Calculator")
                c_calc1, c_calc2 = st.columns([1, 2])
                with c_calc1: desired_profit = st.number_input("Desired Profit ($)", value=50, step=10)
                with c_calc2:
                    target_price_val = 0; move_val = 0
                    if abs(d) > 0.001:
                        move_val = (desired_profit / 100) / abs(d)
                        target_price_val = current_price - move_val if option_type == 'put' else current_price + move_val
                        direction = "DROP -" if option_type == 'put' else "RISE +"
                        color = "#FF4B4B" if option_type == 'put' else "#00FF7F"
                        st.markdown(f"<div style='background-color: #1E3D59; padding: 20px; border-radius: 10px; border-left: 10px solid {color}; color: white; margin-bottom: 20px;'><h4 style='margin:0; color: white;'>Target Stock Price: <b>${target_price_val:.2f}</b></h4><p style='margin:0; color: white;'>Stock needs to {direction}${move_val:.2f}</p></div>", unsafe_allow_html=True)
                st.subheader("🗓️ Holiday Decay Calculator")
                holidays = st.number_input("Days Closed", 1)
                decay_loss_val = abs(t) * holidays * 100
                st.markdown(f"<div style='background-color: #330000; padding: 20px; border-radius: 10px; border-left: 10px solid #FF4B4B; color: white;'><h4 style='margin:0; color: white;'>Estimated Loss: <b>${decay_loss_val:.2f} per contract</b></h4><p style='margin:0; color: white;'>While you sleep...</p></div>", unsafe_allow_html=True)

            with tabs[6]: st.metric("Max Pain", f"${max_pain_val:.2f}")
            with tabs[7]:
                st.header("Latest News (Google News)")
                if news_data:
                    try:
                        for item in news_data[:3]: 
                            st.markdown(f"**{item.get('publisher','')}** - [{item.get('title','')}]({item.get('link','')})  \n*{item.get('published','')}*")
                            st.markdown("---")
                    except: 
                        st.write("News unavailable.")
                else: 
                    st.write("No news found.")

            with tabs[8]: # AI
                st.header("🤖 AI Chart Analyst")
                if "ai_result" not in st.session_state: st.session_state["ai_result"] = ""
                up_files = st.file_uploader("Upload Charts", type=["jpg", "png"], accept_multiple_files=True)
                if up_files and st.button("Analyze Images"):
                    if "api_keys" in st.secrets:
                        genai.configure(api_key=st.secrets["api_keys"]["gemini"])
                        try:
                            model = genai.GenerativeModel("models/gemini-2.0-flash-exp")
                            content = ["Analyze for walls and traps."] + [Image.open(f) for f in up_files]
                            with st.spinner("Analyzing..."):
                                st.session_state["ai_result"] = model.generate_content(content).text; st.rerun()
                        except Exception as e: st.error(str(e))
                if st.session_state["ai_result"]: st.write(st.session_state["ai_result"])

            with tabs[9]: # Strategy
                st.header("💬 Strategy Engine")
                if "strat_log" not in st.session_state: st.session_state["strat_log"] = ""
                q = st.text_input("Ask a strategy question:")
                if q and st.button("Ask AI"):
                     if "api_keys" in st.secrets:
                        genai.configure(api_key=st.secrets["api_keys"]["gemini"])
                        model = genai.GenerativeModel("models/gemini-2.0-flash-exp")
                        resp = model.generate_content(f"Context: {ticker} {option_type.upper()} ${strike_price}. Question: {q}").text
                        st.session_state["strat_log"] = f"Q: {q}\nA: {resp}"; st.write(resp)

            with tabs[10]: # Sim
                st.header("🔮 Future Simulator")
                st.plotly_chart(plot_simulation_interactive(current_price, strike_price, days_left, contract_iv, option_type, purchase_price=theo_price), use_container_width=True)

            with tabs[11]: # Flow
                st.header("🌊 Market Flow")
                st.plotly_chart(plot_flow_battle_interactive(calls, puts, strike_price), use_container_width=True)

            with tabs[12]: # Scanner
                st.header("🔍 ATM Options Scanner")
                
                # --- NEW: MANUAL ENTRY SECTION ---
                col_scan1, col_scan2 = st.columns([3, 1])
                with col_scan1:
                    manual_tickers = st.text_input("Enter Tickers Manually (comma separated, e.g., TSLA, NVDA)", key="manual_ticker_input")
                with col_scan2:
                    st.write("") # Spacer
                    st.write("") # Spacer
                    if st.button("🚀 Scan Manual"):
                        if manual_tickers:
                            tickers = [t.strip().upper() for t in manual_tickers.split(',') if t.strip()]
                            st.session_state["scan_results"] = run_scan(tickers)

                st.markdown("--- OR ---")

                # --- EXISTING: FILE UPLOAD SECTION ---
                up_xl = st.file_uploader("Upload Excel List", type=['xlsx'])
                if up_xl and st.button("🚀 Scan File"):
                    df_input = pd.read_excel(up_xl)
                    tickers = df_input.iloc[:, 0].dropna().astype(str).tolist()
                    st.session_state["scan_results"] = run_scan(tickers)
                
                # --- RESULTS ---
                if "scan_results" in st.session_state:
                    st.dataframe(st.session_state["scan_results"], use_container_width=True)
                    buffer = BytesIO()
                    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer: st.session_state["scan_results"].to_excel(writer, index=False)
                    st.download_button("📥 Download Excel", buffer.getvalue(), "ATM_Scan.xlsx", "application/vnd.ms-excel")
            
            # --- REPORT EXPORT ---
            st.sidebar.markdown("---")
            data_pack = {'ticker': ticker, 'type': option_type, 'price': current_price, 'strike': strike_price, 'exp': selected_date, 'iv': contract_iv, 'daily_move': daily_move, 'volume': info.get('volume', 0), 'max_pain': max_pain_val, 'delta': d, 'gamma': g, 'theta': t, 'profit_goal': desired_profit, 'profit_price': target_price_val, 'holidays': holidays, 'decay_loss': decay_loss_val, 'ai_text': st.session_state.get("ai_result", ""), 'strat_log': st.session_state.get("strat_log", ""), 'news': news_data, 'scan': st.session_state.get("scan_results", None)}
            report_file = generate_full_dossier(data_pack)
            st.sidebar.download_button("📄 Download Full Dossier", report_file.getvalue(), f"{ticker}_{option_type.upper()}_Full_Dossier.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        except Exception as e:
            st.error(f"Waiting for data... ({e})")

# =========================================================
# PAGE 2: PORTFOLIO TRACKER (Separate View)
# =========================================================
elif app_mode == "Portfolio Tracker 📒":
    st.title("📒 Live Portfolio Tracker")
    st.caption("Add trades below. Double-click cells to edit manually.")

    with st.expander("➕ Add New Trade", expanded=True):
        c_pf1, c_pf2, c_pf3 = st.columns(3)
        with c_pf1:
            pf_ticker = st.text_input("Ticker", value="TSLA").upper()
            pf_type = st.selectbox("Type", ["Option", "Stock"])
        with c_pf2:
            if pf_type == "Option":
                pf_strike = st.text_input("Strike", value="440")
                pf_opt_type = st.selectbox("Call/Put", ["Call", "Put"])
                pf_exp = st.text_input("Exp Date", value="Jan 9")
                details = f"{pf_exp} {pf_strike} {pf_opt_type}"
            else:
                details = "Shares"
            pf_qty = st.number_input("Qty", 1, step=1)
        with c_pf3:
            pf_buy_price = st.number_input("Bought Price", value=1.00)
            pf_curr_price = st.number_input("Current Price (Manual)", value=1.00)

        if st.button("Add Trade to List"):
            mult = 100 if pf_type == "Option" else 1
            cost = pf_buy_price * pf_qty * mult
            curr_val = pf_curr_price * pf_qty * mult
            pl = curr_val - cost
            new_row = {"Ticker": pf_ticker, "Type": pf_type, "Details": details, "Qty": pf_qty, "Bought Price": pf_buy_price, "Current Price": pf_curr_price, "Total Cost": cost, "Current Value": curr_val, "Total P/L": pl}
            st.session_state["portfolio"] = pd.concat([st.session_state["portfolio"], pd.DataFrame([new_row])], ignore_index=True)
            st.rerun()

    st.markdown("### 📋 Your Trades (Editable)")
    
    # 2. EDITABLE DATA TABLE (FIX)
    edited_df = st.data_editor(
        st.session_state["portfolio"],
        num_rows="dynamic",
        column_config={
            "Bought Price": st.column_config.NumberColumn(format="$%.2f"),
            "Current Price": st.column_config.NumberColumn(format="$%.2f", help="Update this manually"),
            "Total Cost": st.column_config.NumberColumn(format="$%.2f", disabled=True),
            "Current Value": st.column_config.NumberColumn(format="$%.2f", disabled=True),
            "Total P/L": st.column_config.NumberColumn(format="$%.2f", disabled=True),
        },
        use_container_width=True,
        key="portfolio_editor"
    )

    if not edited_df.equals(st.session_state["portfolio"]):
        recalced_df = recalculate_portfolio(edited_df)
        st.session_state["portfolio"] = recalced_df
        st.rerun()

    if not st.session_state["portfolio"].empty:
        total_pl = st.session_state["portfolio"]['Total P/L'].sum()
        st.markdown(f"### Total Portfolio P/L: :{'green' if total_pl >= 0 else 'red'}[${total_pl:,.2f}]")

